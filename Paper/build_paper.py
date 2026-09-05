# -*- coding: utf-8 -*-
"""Paper/build_paper.py — 국문/영문 논문 원고(.docx) 생성.

모든 본문은 (ko, en) 쌍으로 한 곳에 정의하고, 같은 그림/표를 두 문서가 공유한다.
수치는 전부 이 저장소의 실측 결과(JSON/CSV/보고서)에서 가져온 값이다.

실행: C:/ml_env/Scripts/python Paper/build_paper.py
"""
import os
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.stdout.reconfigure(encoding="utf-8")
HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, "figures")

# ── 문서 스타일 ──────────────────────────────────────────────────────────────
KO_FONT, EN_FONT = "맑은 고딕", "Times New Roman"


def new_doc(lang):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    for a in ("top_margin", "bottom_margin"):
        setattr(sec, a, Cm(2.2))
    for a in ("left_margin", "right_margin"):
        setattr(sec, a, Cm(2.3))
    st = doc.styles["Normal"]
    st.font.name = KO_FONT if lang == "ko" else EN_FONT
    st.font.size = Pt(10.5 if lang == "ko" else 10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), KO_FONT)
    pf = st.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.35 if lang == "ko" else 1.25
    return doc


def _set_font(run, lang, size=None, bold=False, italic=False, color=None):
    run.font.name = KO_FONT if lang == "ko" else EN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), KO_FONT)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, lang, text, size=None, bold=False, italic=False, align=None,
         space_before=0, space_after=6, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align == "c":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "j":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    _set_font(r, lang, size, bold, italic, color)
    return p


def heading(doc, lang, text, level=1):
    sizes = {1: 13, 2: 11.5, 3: 10.5}
    para(doc, lang, text, size=sizes[level], bold=True,
         space_before=14 if level == 1 else 10, space_after=4)


def figure(doc, lang, fname, caption, width_cm=16.0):
    path = os.path.join(FIG, fname)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Cm(width_cm))
    para(doc, lang, caption, size=8.5, align="c", space_after=10)


def table(doc, lang, header, rows, caption=None, widths=None, size=8.5):
    if caption:
        para(doc, lang, caption, size=8.5, bold=True, space_before=8, space_after=3)
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        _set_font(r, lang, size, bold=True)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            _set_font(r, lang, size)
            if i > 0:
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    para(doc, lang, "", size=4, space_after=4)
    return t


def bullets(doc, lang, items, size=None):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(it)
        _set_font(r, lang, size)


def eq(doc, lang, text):
    para(doc, lang, text, size=10, italic=False, align="c",
         space_before=4, space_after=6)


# ── 실행 ────────────────────────────────────────────────────────────────────
def main():
    import content
    import supplementary
    out = []
    for lang, fname in [("ko", "논문_국문_단일전극_멀티모달센서.docx"),
                        ("en", "Paper_EN_SingleElectrode_MultimodalSensor.docx")]:
        doc = new_doc(lang)
        content.build(doc, lang)
        supplementary.build(doc, lang)
        path = os.path.join(HERE, fname)
        doc.save(path)
        out.append(path)
        print("saved:", path)
    return out


if __name__ == "__main__":
    main()
